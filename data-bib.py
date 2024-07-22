import pandas as pd
from docx import Document

# Load the CSV file
csv_file_path = '/Users/radeksabatka/Library/Mobile Documents/com~apple~CloudDocs/Academics/IR499 Dissertation/Python/nuclear-nationalism-china/data/dataset.csv'
df = pd.read_csv(csv_file_path)

# Function to generate bibliography entry in Chicago style
def generate_bibliography(row):
    speaker = row['speaker_name']
    title_chinese = row['title_chinese']
    title_english = row['title_english']
    date = row['date']
    location = row['location']
    context = row['context']
    source = row['url']
    bibliography_entry = f"{speaker}. \"{title_chinese}\" [{title_english}]. Speech delivered at the {context}, {location}, {date}. {source}."
    return bibliography_entry

# Apply the function to each row and store the results
df['Bibliography Entry'] = df.apply(generate_bibliography, axis=1)

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Bibliography', level=1)

# Add each bibliography entry to the document
for entry in df['Bibliography Entry']:
    doc.add_paragraph(entry)

# Save the document to a .docx file
output_file_path = '/Users/radeksabatka/Library/Mobile Documents/com~apple~CloudDocs/Academics/IR499 Dissertation/Python/outputs/bibliography.docx'
doc.save(output_file_path)

print(f"Bibliography saved to '{output_file_path}'.")